"""Génération de l'Acte d'Engagement (ATTRI1) au format Word (.docx).

Produit un Acte d'Engagement conforme au modèle ATTRI1 officiel, adapté
pour les marchés publics de services informatiques, pré-rempli avec les
informations Fenrir IT et les données de l'offre BOAMP.
"""

from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

from entreprise import ENTREPRISE as E


# ── Helper functions (same pattern as dc1_generator.py) ──────────────────────


def _set_cell_shading(cell, color: str) -> None:
    """Set cell background color."""
    from docx.oxml.ns import qn
    from lxml import etree

    shading = etree.SubElement(
        cell._element.get_or_add_tcPr(), qn("w:shd")
    )
    shading.set(qn("w:fill"), color)
    shading.set(qn("w:val"), "clear")


def _add_heading_cell(table, row: int, col: int, text: str) -> None:
    """Add a blue-shaded heading cell."""
    cell = table.cell(row, col)
    cell.text = text
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(11)
    _set_cell_shading(cell, "DAEEF3")


def _add_section_heading(doc: Document, text: str) -> None:
    """Add a blue-background section heading like the official ATTRI1."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = True
    run.underline = True
    run.font.size = Pt(11)
    _set_cell_shading(cell, "DAEEF3")


# ── Main generator ───────────────────────────────────────────────────────────


def generate_ae(notice: dict[str, Any], dest_path: Path) -> Path:
    """Generate an Acte d'Engagement Word document for a BOAMP notice.

    Args:
        notice: BOAMP notice data dict.
        dest_path: Path to save the .docx file.

    Returns:
        Path to the generated file.
    """
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)

    # Extract notice data
    objet = notice.get("objet", "[OBJET DU MARCH\u00c9]")
    acheteur = notice.get("nomacheteur", "[POUVOIR ADJUDICATEUR]")
    idweb = notice.get("idweb", "[REF]")
    type_marche = notice.get("type_marche", "SERVICES")

    # ══════════════════════════════════════════════════════════════════════════
    # Title page
    # ══════════════════════════════════════════════════════════════════════════
    title_p1 = doc.add_paragraph()
    title_p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p1.add_run("MARCH\u00c9 PUBLIC DE SERVICES")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()

    title_p2 = doc.add_paragraph()
    title_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p2.add_run("ACTE D'ENGAGEMENT")
    run.bold = True
    run.font.size = Pt(18)

    doc.add_paragraph()
    doc.add_paragraph()

    # ── Page 1 info block ──
    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = "Table Grid"
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    info_fields = [
        ("Pouvoir adjudicateur", acheteur),
        ("Objet du march\u00e9", objet),
        ("R\u00e9f\u00e9rence BOAMP", idweb),
        ("Type de march\u00e9", type_marche),
    ]

    for i, (label, value) in enumerate(info_fields):
        _set_cell_shading(info_table.cell(i, 0), "DAEEF3")
        label_p = info_table.cell(i, 0).paragraphs[0]
        label_run = label_p.add_run(label)
        label_run.bold = True
        label_run.font.size = Pt(10)

        val_p = info_table.cell(i, 1).paragraphs[0]
        val_run = val_p.add_run(value)
        val_run.font.size = Pt(10)

    for row in info_table.rows:
        row.cells[0].width = Cm(5)
        row.cells[1].width = Cm(12)

    doc.add_paragraph()
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # Article 1 - Contractant
    # ══════════════════════════════════════════════════════════════════════════
    _add_section_heading(doc, "Article 1 - Identification du contractant")

    doc.add_paragraph()

    a1_intro = doc.add_paragraph()
    a1_intro.add_run("Je soussign\u00e9,").italic = True

    doc.add_paragraph()

    a1_table = doc.add_table(rows=8, cols=2)
    a1_table.style = "Table Grid"
    a1_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    a1_fields = [
        ("Nom et pr\u00e9nom", f"{E['representant_prenom']} {E['representant_nom']}"),
        ("Agissant en qualit\u00e9 de", E["representant_qualite"]),
        ("D\u00e9nomination sociale", E["raison_sociale"]),
        ("Forme juridique", E["forme_juridique"]),
        ("Capital social", E["capital_social"]),
        (
            "Adresse du si\u00e8ge social",
            f"{E['adresse']}\n{E['code_postal']} {E['ville']}",
        ),
        ("Num\u00e9ro SIRET", E["siret"]),
        ("Adresse \u00e9lectronique / T\u00e9l\u00e9phone", f"{E['email']} / {E['telephone']}"),
    ]

    for i, (label, value) in enumerate(a1_fields):
        _set_cell_shading(a1_table.cell(i, 0), "F2F2F2")
        label_p = a1_table.cell(i, 0).paragraphs[0]
        label_run = label_p.add_run(label)
        label_run.bold = True
        label_run.font.size = Pt(9)

        val_p = a1_table.cell(i, 1).paragraphs[0]
        val_run = val_p.add_run(value)
        val_run.font.size = Pt(10)

    for row in a1_table.rows:
        row.cells[0].width = Cm(6)
        row.cells[1].width = Cm(11)

    doc.add_paragraph()

    a1_engage = doc.add_paragraph()
    run = a1_engage.add_run(
        "agissant au nom et pour le compte de la soci\u00e9t\u00e9 susmentionn\u00e9e, "
        "m'engage, conform\u00e9ment aux stipulations du pr\u00e9sent acte d'engagement, "
        "du cahier des clauses administratives particuli\u00e8res (CCAP) et du cahier "
        "des clauses techniques particuli\u00e8res (CCTP), \u00e0 ex\u00e9cuter les prestations "
        "du march\u00e9 d\u00e9sign\u00e9 ci-dessus."
    )
    run.font.size = Pt(10)

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # Article 2 - Objet et prix
    # ══════════════════════════════════════════════════════════════════════════
    _add_section_heading(doc, "Article 2 - Objet du march\u00e9 et prix")

    doc.add_paragraph()

    a2_objet = doc.add_paragraph()
    a2_objet.add_run("Objet : ").bold = True
    a2_objet.add_run(objet).font.size = Pt(10)

    doc.add_paragraph()

    a2_engage = doc.add_paragraph()
    run = a2_engage.add_run(
        "Le titulaire s'engage \u00e0 ex\u00e9cuter les prestations d\u00e9finies dans le CCAP "
        "et le CCTP pour les prix ci-apr\u00e8s :"
    )
    run.font.size = Pt(10)

    doc.add_paragraph()

    # Price table
    price_table = doc.add_table(rows=4, cols=2)
    price_table.style = "Table Grid"
    price_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    price_labels = [
        ("Montant HT (\u20ac)", ""),
        ("TVA 20 % (\u20ac)", ""),
        ("Montant TTC (\u20ac)", ""),
    ]

    # Header row
    _set_cell_shading(price_table.cell(0, 0), "DAEEF3")
    _set_cell_shading(price_table.cell(0, 1), "DAEEF3")
    h0 = price_table.cell(0, 0).paragraphs[0]
    h0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h0.add_run("D\u00e9signation")
    run.bold = True
    run.font.size = Pt(10)
    h1 = price_table.cell(0, 1).paragraphs[0]
    h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h1.add_run("Montant")
    run.bold = True
    run.font.size = Pt(10)

    for i, (label, value) in enumerate(price_labels, start=1):
        label_p = price_table.cell(i, 0).paragraphs[0]
        label_run = label_p.add_run(label)
        label_run.bold = True
        label_run.font.size = Pt(10)

        val_p = price_table.cell(i, 1).paragraphs[0]
        val_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        val_run = val_p.add_run(value)
        val_run.font.size = Pt(10)

    for row in price_table.rows:
        row.cells[0].width = Cm(8)
        row.cells[1].width = Cm(6)

    doc.add_paragraph()

    a2_validity = doc.add_paragraph()
    a2_validity.add_run("Dur\u00e9e de validit\u00e9 de l'offre : ").bold = True
    a2_validity.add_run("120 jours \u00e0 compter de la date limite de remise des offres.").font.size = Pt(10)

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # Article 3 - Durée
    # ══════════════════════════════════════════════════════════════════════════
    _add_section_heading(doc, "Article 3 - Dur\u00e9e du march\u00e9")

    doc.add_paragraph()

    a3_duree = doc.add_paragraph()
    a3_duree.add_run("Dur\u00e9e du march\u00e9 : ").bold = True
    a3_duree.add_run("________________________________").font.size = Pt(10)

    a3_recond = doc.add_paragraph()
    a3_recond.add_run("Nombre de reconductions possibles : ").bold = True
    a3_recond.add_run("________________________________").font.size = Pt(10)

    a3_debut = doc.add_paragraph()
    a3_debut.add_run("Date pr\u00e9visionnelle de d\u00e9but d'ex\u00e9cution : ").bold = True
    a3_debut.add_run("________________________________").font.size = Pt(10)

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # Article 4 - Paiements
    # ══════════════════════════════════════════════════════════════════════════
    _add_section_heading(doc, "Article 4 - Modalit\u00e9s de r\u00e8glement et de paiement")

    doc.add_paragraph()

    a4_intro = doc.add_paragraph()
    run = a4_intro.add_run(
        "Les sommes dues au titre du pr\u00e9sent march\u00e9 seront r\u00e9gl\u00e9es par virement "
        "bancaire au compte d\u00e9sign\u00e9 ci-dessous :"
    )
    run.font.size = Pt(10)

    doc.add_paragraph()

    bank_table = doc.add_table(rows=4, cols=2)
    bank_table.style = "Table Grid"
    bank_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    bank_fields = [
        ("\u00c9tablissement bancaire", E["banque_etablissement"]),
        ("Titulaire du compte", E["raison_sociale"]),
        ("IBAN", E["banque_iban"]),
        ("BIC", E["banque_bic"]),
    ]

    for i, (label, value) in enumerate(bank_fields):
        _set_cell_shading(bank_table.cell(i, 0), "F2F2F2")
        label_p = bank_table.cell(i, 0).paragraphs[0]
        label_run = label_p.add_run(label)
        label_run.bold = True
        label_run.font.size = Pt(9)

        val_p = bank_table.cell(i, 1).paragraphs[0]
        val_run = val_p.add_run(value)
        val_run.font.size = Pt(10)

    for row in bank_table.rows:
        row.cells[0].width = Cm(5)
        row.cells[1].width = Cm(12)

    doc.add_paragraph()

    a4_delai = doc.add_paragraph()
    a4_delai.add_run("D\u00e9lai global de paiement : ").bold = True
    a4_delai.add_run(
        "30 jours conform\u00e9ment aux dispositions de l'article R. 2192-10 "
        "du Code de la commande publique."
    ).font.size = Pt(10)

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # Article 5 - Signature
    # ══════════════════════════════════════════════════════════════════════════
    _add_section_heading(doc, "Article 5 - Signature de l'acte d'engagement")

    doc.add_paragraph()

    a5_text = doc.add_paragraph()
    run = a5_text.add_run(
        "Le contractant affirme, sous peine de r\u00e9siliation du march\u00e9 \u00e0 ses torts, "
        "ne pas avoir fait l'objet d'une interdiction de concourir et que les "
        "renseignements fournis dans le cadre de cette consultation sont exacts."
    )
    run.font.size = Pt(10)

    doc.add_paragraph()

    sig_table = doc.add_table(rows=2, cols=3)
    sig_table.style = "Table Grid"
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    sig_headers = [
        "Nom, pr\u00e9nom et qualit\u00e9\ndu signataire",
        "Lieu et date",
        "Signature",
    ]
    for i, h in enumerate(sig_headers):
        _set_cell_shading(sig_table.cell(0, i), "DAEEF3")
        p = sig_table.cell(0, i).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)

    # Pre-fill signature row
    sig_table.cell(1, 0).paragraphs[0].add_run(
        f"{E['representant_prenom']} {E['representant_nom']}\n{E['representant_qualite']}"
    ).font.size = Pt(10)

    sig_table.cell(1, 1).paragraphs[0].add_run(
        f"{E['ville']}, le ____/____/________"
    ).font.size = Pt(10)

    sig_table.cell(1, 2).paragraphs[0].add_run(
        "\n\n\n"
    )

    # Set signature row height for space
    for cell in sig_table.rows[1].cells:
        cell.height = Cm(3)

    doc.add_paragraph()

    # ── Footer ──
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_p.add_run(
        f"Acte d'Engagement          {idweb}          "
        f"Page 1 / 1"
    )
    footer_run.font.size = Pt(8)
    footer_run.italic = True

    # Save
    dest_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(dest_path))
    return dest_path
