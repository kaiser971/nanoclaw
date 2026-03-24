"""Génération du Mémoire Technique au format Word (.docx).

Produit un mémoire technique structuré pour les réponses aux appels d'offres
suivant le modèle Odialis, pré-rempli avec les informations Fenrir IT
et les données de l'offre BOAMP.

Offer types supportés : TMA, DEVELOPPEMENT, FORMATION, IA.
"""

from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

from entreprise import ENTREPRISE as E


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _set_cell_shading(cell, color: str) -> None:
    """Set cell background color."""
    from docx.oxml.ns import qn
    from lxml import etree

    shading = etree.SubElement(
        cell._element.get_or_add_tcPr(), qn("w:shd")
    )
    shading.set(qn("w:fill"), color)
    shading.set(qn("w:val"), "clear")


def _set_run_font(run, name: str = "Arial", size: int = 10,
                  bold: bool = False, italic: bool = False,
                  color: RGBColor | None = None) -> None:
    """Apply common font settings to a run."""
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def _add_section_heading(doc: Document, number: int, title: str) -> None:
    """Add a numbered blue-background section heading."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    p = cell.paragraphs[0]
    run = p.add_run(f"{number}. {title}")
    _set_run_font(run, bold=True, size=13, color=RGBColor(0x1F, 0x49, 0x7D))
    _set_cell_shading(cell, "DAEEF3")
    doc.add_paragraph()  # spacing after heading


def _add_table_row(table, row_idx: int, values: list[str],
                   bold: bool = False, header: bool = False) -> None:
    """Fill a table row with values."""
    for col_idx, val in enumerate(values):
        cell = table.cell(row_idx, col_idx)
        p = cell.paragraphs[0]
        run = p.add_run(val)
        _set_run_font(run, size=10, bold=bold)
        if header:
            _set_cell_shading(cell, "DAEEF3")


def _make_grid_table(doc: Document, headers: list[str],
                     rows: list[list[str]]) -> None:
    """Create a Table Grid table with headers and data rows."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _add_table_row(table, 0, headers, bold=True, header=True)
    for i, row_data in enumerate(rows, start=1):
        _add_table_row(table, i, row_data)
    doc.add_paragraph()


def _add_sub_heading(doc: Document, text: str) -> None:
    """Add a bold sub-heading paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run_font(run, bold=True, size=11, color=RGBColor(0x1F, 0x49, 0x7D))
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)


def _add_placeholder(doc: Document, text: str) -> None:
    """Add an italic grey placeholder paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run_font(run, italic=True, size=10, color=RGBColor(0x80, 0x80, 0x80))


# ---------------------------------------------------------------------------
# TOC titles per offer type
# ---------------------------------------------------------------------------

_TOC_ITEMS: dict[str, list[str]] = {
    "TMA": [
        "Présentation de l'entreprise",
        "Compréhension du besoin",
        "Méthodologie TMA proposée",
        "Organisation et gouvernance",
        "Équipe dédiée",
        "Outillage et environnement technique",
        "Engagements de service (SLA)",
        "Gestion de la transition / réversibilité",
        "Références clients",
        "Annexes",
    ],
    "DEVELOPPEMENT": [
        "Présentation de l'entreprise",
        "Compréhension du besoin",
        "Méthodologie de développement web",
        "Organisation projet",
        "Équipe projet",
        "Stack technique et outillage",
        "Engagements qualité",
        "Livraison et garantie",
        "Références clients",
        "Annexes",
    ],
    "FORMATION": [
        "Présentation de l'entreprise",
        "Compréhension du besoin",
        "Ingénierie pédagogique et méthodologie",
        "Organisation et gouvernance",
        "Équipe pédagogique",
        "Outils et plateforme",
        "Indicateurs de performance",
        "Transfert et pérennité",
        "Références clients",
        "Annexes",
    ],
    "IA": [
        "Présentation de l'entreprise",
        "Compréhension du besoin",
        "Méthodologie IA",
        "Gouvernance IA",
        "Équipe IA",
        "Stack technique IA",
        "Engagements performance IA",
        "Réversibilité et éthique",
        "Références clients",
        "Annexes",
    ],
}


# ---------------------------------------------------------------------------
# Type-specific sections (3-8)
# ---------------------------------------------------------------------------

def _sections_tma(doc: Document, notice: dict[str, Any]) -> None:
    """Sections 3-8 for TMA offer type."""

    # ── Section 3 — Méthodologie TMA proposée ──
    _add_section_heading(doc, 3, "Méthodologie TMA proposée")

    _add_sub_heading(doc, "3.1 Maintenance corrective")
    p = doc.add_paragraph()
    run = p.add_run(
        "La maintenance corrective vise à corriger les anomalies détectées "
        "en production ou en pré-production. Chaque incident est qualifié "
        "selon une grille de priorité définie conjointement avec l'acheteur."
    )
    _set_run_font(run, size=10)

    _make_grid_table(
        doc,
        ["Priorité", "Description", "Délai de prise en compte", "Délai de résolution"],
        [
            ["P1 — Bloquant", "Application indisponible, perte de données", "[À compléter]", "[À compléter]"],
            ["P2 — Majeur", "Fonctionnalité critique dégradée", "[À compléter]", "[À compléter]"],
            ["P3 — Mineur", "Dysfonctionnement non bloquant", "[À compléter]", "[À compléter]"],
            ["P4 — Cosmétique", "Défaut visuel, ergonomique", "[À compléter]", "[À compléter]"],
        ],
    )

    _add_sub_heading(doc, "3.2 Maintenance évolutive")
    p = doc.add_paragraph()
    run = p.add_run(
        "La maintenance évolutive couvre les adaptations fonctionnelles et "
        "techniques demandées par l'acheteur. Le processus suit les étapes "
        "suivantes :"
    )
    _set_run_font(run, size=10)

    steps = [
        "Réception et analyse de la demande d'évolution",
        "Chiffrage et proposition de planning",
        "Validation par le comité de pilotage",
        "Spécifications détaillées et conception",
        "Développement et tests unitaires",
        "Recette interne puis recette acheteur",
        "Mise en production et documentation",
    ]
    for i, step in enumerate(steps, start=1):
        p = doc.add_paragraph()
        run = p.add_run(f"  {i}. {step}")
        _set_run_font(run, size=10)

    _add_sub_heading(doc, "3.3 Maintenance préventive")
    _add_placeholder(
        doc,
        "[Décrire les actions de maintenance préventive : veille "
        "technologique, mises à jour de sécurité, optimisation des "
        "performances, revue de code, audits de vulnérabilité, gestion "
        "de l'obsolescence technique.]"
    )

    _add_sub_heading(doc, "3.4 Processus global — Vue synthétique")
    process_lines = [
        "┌──────────────┐    ┌─────────────────┐    ┌───────────────┐",
        "│  RÉCEPTION   │───▶│  QUALIFICATION  │───▶│   AFFECTATION │",
        "│  Demande     │    │  Priorité/Type   │    │   Ressource   │",
        "└──────────────┘    └─────────────────┘    └───────┬───────┘",
        "                                                   │        ",
        "┌──────────────┐    ┌─────────────────┐    ┌───────▼───────┐",
        "│  CLÔTURE     │◀───│  RECETTE        │◀───│   RÉSOLUTION  │",
        "│  Bilan       │    │  Validation      │    │   Dev/Correctif│",
        "└──────────────┘    └─────────────────┘    └───────────────┘",
    ]
    for line in process_lines:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(8)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)

    doc.add_page_break()

    # ── Section 4 — Organisation et gouvernance ──
    _add_section_heading(doc, 4, "Organisation et gouvernance")

    _add_sub_heading(doc, "4.1 Instances de pilotage")
    _make_grid_table(
        doc,
        ["Instance", "Fréquence", "Participants", "Objectif"],
        [
            [
                "Comité de pilotage",
                "Trimestriel",
                "Direction Fenrir IT, DSI acheteur",
                "Bilan global, arbitrages stratégiques, feuille de route",
            ],
            [
                "Comité de suivi",
                "Mensuel",
                "Chef de projet TMA, Référent acheteur",
                "Suivi des indicateurs, planning, risques",
            ],
            [
                "Point opérationnel",
                "Hebdomadaire",
                "Équipe TMA, Référent technique acheteur",
                "Avancement tickets, priorisation, points bloquants",
            ],
        ],
    )

    _add_sub_heading(doc, "4.2 Reporting et livrables de pilotage")
    reporting_items = [
        "Tableau de bord mensuel (indicateurs SLA, volumétrie, tendances)",
        "Compte-rendu de chaque comité",
        "Rapport d'activité trimestriel",
        "Plan d'amélioration continue",
    ]
    for item in reporting_items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        _set_run_font(run, size=10)

    doc.add_page_break()

    # ── Section 5 — Équipe dédiée ──
    _add_section_heading(doc, 5, "Équipe dédiée")

    _add_sub_heading(doc, "5.1 Organigramme projet")
    _add_placeholder(
        doc,
        "[Insérer ici l'organigramme de l'équipe projet. Présenter la "
        "chaîne hiérarchique et fonctionnelle, du directeur de projet "
        "jusqu'aux développeurs.]"
    )

    _add_sub_heading(doc, "5.2 Profils mobilisés")
    _make_grid_table(
        doc,
        ["Rôle", "Nom", "Expérience", "Taux d'occupation"],
        [
            ["Directeur de projet", "[Nom]", "[X ans]", "[X %]"],
            ["Chef de projet TMA", "[Nom]", "[X ans]", "[X %]"],
            ["Architecte technique", "[Nom]", "[X ans]", "[X %]"],
            ["Développeur senior", "[Nom]", "[X ans]", "[X %]"],
            ["Développeur", "[Nom]", "[X ans]", "[X %]"],
            ["Analyste / QA", "[Nom]", "[X ans]", "[X %]"],
        ],
    )

    _add_placeholder(
        doc,
        "[Les CV détaillés de chaque intervenant sont fournis en annexe.]"
    )

    doc.add_page_break()

    # ── Section 6 — Outillage et environnement technique ──
    _add_section_heading(doc, 6, "Outillage et environnement technique")

    _make_grid_table(
        doc,
        ["Fonction", "Outil proposé"],
        [
            ["Gestion de tickets / ITSM", "JIRA Service Management / GLPI"],
            ["Gestion de projet", "JIRA / Azure DevOps"],
            ["Gestion de configuration (SCM)", "Git (GitLab / GitHub)"],
            ["Intégration continue (CI/CD)", "GitLab CI / Jenkins / GitHub Actions"],
            ["Gestion documentaire", "Confluence / SharePoint"],
            ["Supervision applicative", "Grafana / Prometheus / Zabbix"],
            ["Tests automatisés", "Selenium / Cypress / JUnit"],
            ["Sécurité / Analyse de code", "SonarQube / Snyk"],
            ["Communication", "Microsoft Teams / Slack"],
        ],
    )

    _add_placeholder(
        doc,
        "[Adapter la liste d'outils en fonction des exigences du CCTP "
        "et de l'environnement technique de l'acheteur.]"
    )

    doc.add_page_break()

    # ── Section 7 — Engagements de service (SLA) ──
    _add_section_heading(doc, 7, "Engagements de service (SLA)")

    _make_grid_table(
        doc,
        ["Indicateur", "Cible", "Pénalité"],
        [
            [
                "Taux de disponibilité",
                "99,5 %",
                "[X % du forfait mensuel par tranche de 0,1 % en dessous]",
            ],
            [
                "Délai de prise en compte P1",
                "≤ 1 heure (HO)",
                "[À définir]",
            ],
            [
                "Délai de résolution P1",
                "≤ 4 heures (HO)",
                "[À définir]",
            ],
            [
                "Délai de résolution P2",
                "≤ 1 jour ouvré",
                "[À définir]",
            ],
            [
                "Taux de résolution au 1er contact",
                "≥ 80 %",
                "[À définir]",
            ],
            [
                "Respect des délais évolutions",
                "≥ 90 % des jalons respectés",
                "[À définir]",
            ],
            [
                "Satisfaction utilisateur",
                "≥ 4/5 (enquête trimestrielle)",
                "[À définir]",
            ],
        ],
    )

    _add_placeholder(
        doc,
        "[Adapter les cibles et pénalités en fonction des exigences "
        "du CCAP et du CCTP.]"
    )

    doc.add_page_break()

    # ── Section 8 — Gestion de la transition / réversibilité ──
    _add_section_heading(doc, 8, "Gestion de la transition / réversibilité")

    _add_sub_heading(doc, "8.1 Plan de transition (phase d'initialisation)")
    _make_grid_table(
        doc,
        ["Phase", "Durée estimée", "Activités clés", "Livrables"],
        [
            [
                "1. Cadrage",
                "2 semaines",
                "Réunion de lancement, recueil documentaire, "
                "identification des interlocuteurs",
                "Plan de transition, matrice RACI",
            ],
            [
                "2. Montée en compétence",
                "4 semaines",
                "Étude du patrimoine applicatif, formation, "
                "accès aux environnements",
                "Dossier d'architecture, base de connaissances",
            ],
            [
                "3. Run accompagné",
                "4 semaines",
                "Prise en charge progressive des tickets avec "
                "accompagnement du titulaire sortant",
                "Tableau de bord de transition",
            ],
            [
                "4. Autonomie",
                "2 semaines",
                "Gestion autonome, clôture de la transition",
                "PV de fin de transition, bilan",
            ],
        ],
    )

    _add_sub_heading(doc, "8.2 Réversibilité")
    _add_placeholder(
        doc,
        "[Décrire le plan de réversibilité : conditions de transfert "
        "de compétences vers un nouveau titulaire ou vers l'acheteur, "
        "restitution de la documentation, des codes sources, des données, "
        "des accès. Préciser le délai de préavis et la durée de la phase "
        "de réversibilité.]"
    )

    doc.add_page_break()


def _sections_dev(doc: Document, notice: dict[str, Any]) -> None:
    """Sections 3-8 for DEVELOPPEMENT offer type."""

    # ── Section 3 — Méthodologie de développement web ──
    _add_section_heading(doc, 3, "Méthodologie de développement web")

    _add_sub_heading(doc, "3.1 Approche Agile / Scrum")
    p = doc.add_paragraph()
    run = p.add_run(
        "Notre méthodologie de développement repose sur le framework Agile/Scrum, "
        "garantissant une livraison itérative et incrémentale. Le processus "
        "s'articule autour des phases suivantes : discovery, UX/UI, "
        "développement, tests, déploiement et VSR (Vérification de Service Régulier)."
    )
    _set_run_font(run, size=10)

    _add_sub_heading(doc, "3.2 Cycle de sprint")
    _make_grid_table(
        doc,
        ["Cérémonie", "Fréquence", "Durée", "Objectif"],
        [
            ["Sprint planning", "Début de sprint", "2-4h", "Planification du backlog de sprint"],
            ["Daily standup", "Quotidien", "15 min", "Synchronisation équipe, levée des blocages"],
            ["Sprint review", "Fin de sprint (bi-hebdo)", "1-2h", "Démonstration des livrables, feedback client"],
            ["Rétrospective", "Fin de sprint", "1h", "Amélioration continue du processus"],
        ],
    )

    _add_sub_heading(doc, "3.3 Phases du projet")
    phases = [
        "Discovery : ateliers de cadrage, recueil des besoins, user stories",
        "UX/UI : wireframes, maquettes, prototypage, tests utilisateurs",
        "Développement : sprints itératifs, revue de code, intégration continue",
        "Tests : tests unitaires, intégration, end-to-end, recette",
        "Déploiement : mise en production, monitoring, hotfix",
        "VSR : vérification de service régulier, stabilisation",
    ]
    for phase in phases:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(phase)
        _set_run_font(run, size=10)

    doc.add_page_break()

    # ── Section 4 — Organisation projet ──
    _add_section_heading(doc, 4, "Organisation projet")

    _add_sub_heading(doc, "4.1 Gouvernance")
    _make_grid_table(
        doc,
        ["Instance", "Fréquence", "Participants", "Objectif"],
        [
            [
                "Comité de pilotage",
                "Trimestriel",
                "Direction Fenrir IT, DSI acheteur",
                "Arbitrages stratégiques, feuille de route, budget",
            ],
            [
                "Sprint review",
                "Bi-hebdomadaire",
                "Équipe projet, Product Owner acheteur",
                "Démonstration des livrables, validation, priorisation",
            ],
            [
                "Daily standup",
                "Quotidien",
                "Équipe de développement",
                "Synchronisation, levée des blocages",
            ],
        ],
    )

    _add_sub_heading(doc, "4.2 Reporting")
    reporting_items = [
        "Burndown chart et vélocité par sprint",
        "Tableau de bord projet (avancement, risques, budget)",
        "Compte-rendu de sprint review",
        "Rapport d'avancement mensuel",
    ]
    for item in reporting_items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        _set_run_font(run, size=10)

    doc.add_page_break()

    # ── Section 5 — Équipe projet ──
    _add_section_heading(doc, 5, "Équipe projet")

    _add_sub_heading(doc, "5.1 Profils mobilisés")
    _make_grid_table(
        doc,
        ["Rôle", "Nom", "Expérience", "Taux d'occupation"],
        [
            ["Chef de projet", "[Nom]", "[X ans]", "[X %]"],
            ["UX Designer", "[Nom]", "[X ans]", "[X %]"],
            ["Développeur fullstack senior", "[Nom]", "[X ans]", "[X %]"],
            ["Développeur frontend", "[Nom]", "[X ans]", "[X %]"],
            ["Développeur backend", "[Nom]", "[X ans]", "[X %]"],
            ["DevOps", "[Nom]", "[X ans]", "[X %]"],
            ["QA / Testeur", "[Nom]", "[X ans]", "[X %]"],
        ],
    )

    _add_placeholder(
        doc,
        "[Les CV détaillés de chaque intervenant sont fournis en annexe.]"
    )

    doc.add_page_break()

    # ── Section 6 — Stack technique et outillage ──
    _add_section_heading(doc, 6, "Stack technique et outillage")

    _make_grid_table(
        doc,
        ["Domaine", "Technologies / Outils"],
        [
            ["Frontend", "React / Vue.js / Angular"],
            ["Backend", "Node.js / Python / PHP"],
            ["Base de données", "PostgreSQL / MySQL"],
            ["CI/CD", "GitLab CI / GitHub Actions"],
            ["Design", "Figma"],
            ["Gestion de projet", "Jira"],
        ],
    )

    _add_placeholder(
        doc,
        "[Adapter la stack technique en fonction des exigences du CCTP "
        "et de l'environnement technique de l'acheteur.]"
    )

    doc.add_page_break()

    # ── Section 7 — Engagements qualité ──
    _add_section_heading(doc, 7, "Engagements qualité")

    _make_grid_table(
        doc,
        ["Indicateur", "Cible", "Mesure"],
        [
            ["Couverture de tests", "> 80 %", "Rapport de couverture automatisé"],
            ["Vélocité stable", "Écart < 15 % entre sprints", "Burndown chart"],
            ["Délai correction bugs critiques", "< 24h", "Suivi JIRA"],
            ["Conformité RGAA accessibilité", "Niveau AA", "Audit accessibilité"],
        ],
    )

    _add_placeholder(
        doc,
        "[Adapter les indicateurs et cibles en fonction des exigences "
        "du CCAP et du CCTP.]"
    )

    doc.add_page_break()

    # ── Section 8 — Livraison et garantie ──
    _add_section_heading(doc, 8, "Livraison et garantie")

    _add_sub_heading(doc, "8.1 Phases de livraison")
    p = doc.add_paragraph()
    run = p.add_run(
        "La livraison s'effectue de manière itérative à chaque fin de sprint, "
        "avec une mise en production progressive. Chaque livraison inclut : "
        "le code source versionné, la documentation technique, les résultats "
        "de tests, et le procès-verbal de recette."
    )
    _set_run_font(run, size=10)

    _add_sub_heading(doc, "8.2 Période de garantie")
    p = doc.add_paragraph()
    run = p.add_run(
        "Une période de garantie est prévue après la livraison finale. "
        "Durant cette période, toute anomalie liée aux développements "
        "réalisés sera corrigée sans surcoût. La durée et les conditions "
        "de la garantie sont définies conformément au CCAP."
    )
    _set_run_font(run, size=10)

    _add_placeholder(
        doc,
        "[Préciser la durée de garantie proposée, les conditions "
        "d'intervention, et le périmètre couvert.]"
    )

    doc.add_page_break()


def _sections_formation(doc: Document, notice: dict[str, Any]) -> None:
    """Sections 3-8 for FORMATION offer type."""

    # ── Section 3 — Ingénierie pédagogique et méthodologie ──
    _add_section_heading(doc, 3, "Ingénierie pédagogique et méthodologie")

    _add_sub_heading(doc, "3.1 Analyse des besoins")
    p = doc.add_paragraph()
    run = p.add_run(
        "L'ingénierie pédagogique débute par une analyse approfondie des "
        "besoins de formation : public cible, prérequis, objectifs "
        "pédagogiques, contraintes organisationnelles et techniques."
    )
    _set_run_font(run, size=10)

    _add_sub_heading(doc, "3.2 Conception de contenu — Modèle ADDIE")
    p = doc.add_paragraph()
    run = p.add_run(
        "La conception des contenus suit le modèle ADDIE, structuré en "
        "cinq phases :"
    )
    _set_run_font(run, size=10)

    _make_grid_table(
        doc,
        ["Phase", "Description", "Livrables"],
        [
            ["Analyse", "Recueil des besoins, profil des apprenants, objectifs", "Cahier des charges pédagogique"],
            ["Design", "Scénarisation, architecture de formation, modalités", "Storyboard, plan de formation"],
            ["Développement", "Création des supports, modules e-learning, exercices", "Contenus multimédia, quiz"],
            ["Implémentation", "Déploiement sur LMS, formation des formateurs", "Plateforme configurée, accès"],
            ["Évaluation", "Mesure des acquis, satisfaction, amélioration continue", "Rapports d'évaluation, KPIs"],
        ],
    )

    _add_sub_heading(doc, "3.3 Approche pédagogique")
    _add_placeholder(
        doc,
        "[Décrire l'approche pédagogique : blended learning, "
        "microlearning, gamification, classes virtuelles, ateliers "
        "pratiques, études de cas. Adapter en fonction du public cible "
        "et des objectifs du marché.]"
    )

    doc.add_page_break()

    # ── Section 4 — Organisation et gouvernance ──
    _add_section_heading(doc, 4, "Organisation et gouvernance")

    _add_sub_heading(doc, "4.1 Instances de pilotage")
    _make_grid_table(
        doc,
        ["Instance", "Fréquence", "Participants", "Objectif"],
        [
            [
                "Comité pédagogique",
                "Mensuel",
                "Direction projet, référent formation acheteur",
                "Validation des contenus, arbitrages pédagogiques",
            ],
            [
                "Point de suivi",
                "Hebdomadaire",
                "Chef de projet, ingénieur pédagogique",
                "Avancement production, planning, points bloquants",
            ],
            [
                "Revue de contenu",
                "À chaque livraison de module",
                "Équipe pédagogique, experts métier acheteur",
                "Validation du contenu, corrections",
            ],
        ],
    )

    _add_sub_heading(doc, "4.2 Reporting")
    reporting_items = [
        "Tableau de bord de production des contenus",
        "Rapport de suivi des sessions de formation",
        "Synthèse des évaluations stagiaires",
        "Plan d'amélioration continue",
    ]
    for item in reporting_items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        _set_run_font(run, size=10)

    doc.add_page_break()

    # ── Section 5 — Équipe pédagogique ──
    _add_section_heading(doc, 5, "Équipe pédagogique")

    _add_sub_heading(doc, "5.1 Profils mobilisés")
    _make_grid_table(
        doc,
        ["Rôle", "Nom", "Expérience", "Taux d'occupation"],
        [
            ["Directeur de projet", "[Nom]", "[X ans]", "[X %]"],
            ["Ingénieur pédagogique", "[Nom]", "[X ans]", "[X %]"],
            ["Formateur expert web", "[Nom]", "[X ans]", "[X %]"],
            ["Développeur e-learning", "[Nom]", "[X ans]", "[X %]"],
            ["Graphiste / motion designer", "[Nom]", "[X ans]", "[X %]"],
            ["Intégrateur LMS", "[Nom]", "[X ans]", "[X %]"],
        ],
    )

    _add_placeholder(
        doc,
        "[Les CV détaillés de chaque intervenant sont fournis en annexe.]"
    )

    doc.add_page_break()

    # ── Section 6 — Outils et plateforme ──
    _add_section_heading(doc, 6, "Outils et plateforme")

    _make_grid_table(
        doc,
        ["Domaine", "Outils proposés"],
        [
            ["LMS", "Moodle / 360Learning / Talentsoft"],
            ["Authoring", "Articulate / iSpring"],
            ["Visioconférence", "Teams / Zoom"],
            ["Évaluation", "Quiz, certification"],
        ],
    )

    _add_placeholder(
        doc,
        "[Adapter la liste d'outils en fonction des exigences du CCTP "
        "et de l'environnement technique de l'acheteur.]"
    )

    doc.add_page_break()

    # ── Section 7 — Indicateurs de performance ──
    _add_section_heading(doc, 7, "Indicateurs de performance")

    _make_grid_table(
        doc,
        ["Indicateur", "Cible", "Mesure"],
        [
            ["Taux de complétion", "> 85 %", "Statistiques LMS"],
            ["Satisfaction stagiaires", "> 4/5", "Enquête post-formation"],
            ["Taux de réussite certification", "> 90 %", "Résultats certification"],
            ["NPS (Net Promoter Score)", "[À définir]", "Enquête trimestrielle"],
        ],
    )

    _add_placeholder(
        doc,
        "[Adapter les indicateurs et cibles en fonction des exigences "
        "du CCAP et du CCTP.]"
    )

    doc.add_page_break()

    # ── Section 8 — Transfert et pérennité ──
    _add_section_heading(doc, 8, "Transfert et pérennité")

    _add_sub_heading(doc, "8.1 Transfert des contenus")
    p = doc.add_paragraph()
    run = p.add_run(
        "L'ensemble des contenus pédagogiques produits (modules e-learning, "
        "supports de formation, quiz, vidéos) sont livrés dans des formats "
        "standards et ouverts, permettant leur réutilisation et mise à jour "
        "par l'acheteur."
    )
    _set_run_font(run, size=10)

    _add_sub_heading(doc, "8.2 Formation des administrateurs")
    p = doc.add_paragraph()
    run = p.add_run(
        "Une formation spécifique est dispensée aux administrateurs de la "
        "plateforme LMS afin de garantir leur autonomie dans la gestion "
        "quotidienne : gestion des utilisateurs, suivi des parcours, "
        "extraction des statistiques."
    )
    _set_run_font(run, size=10)

    _add_sub_heading(doc, "8.3 Processus de mise à jour")
    _add_placeholder(
        doc,
        "[Décrire le processus de mise à jour des contenus : fréquence "
        "de révision, modalités de mise à jour, gestion des versions, "
        "procédure de validation.]"
    )

    doc.add_page_break()


def _sections_ia(doc: Document, notice: dict[str, Any]) -> None:
    """Sections 3-8 for IA offer type."""

    # ── Section 3 — Méthodologie IA ──
    _add_section_heading(doc, 3, "Méthodologie IA")

    _add_sub_heading(doc, "3.1 Phases du projet IA")
    p = doc.add_paragraph()
    run = p.add_run(
        "Notre méthodologie IA s'appuie sur une approche MLOps structurée, "
        "garantissant la reproductibilité, la traçabilité et la qualité "
        "des modèles déployés en production."
    )
    _set_run_font(run, size=10)

    _make_grid_table(
        doc,
        ["Phase", "Description", "Livrables"],
        [
            ["Audit données", "Qualification des sources de données, analyse de qualité, identification des biais", "Rapport d'audit données"],
            ["Cadrage / PoC", "Définition du périmètre, preuve de concept, validation de faisabilité", "PoC validé, spécifications"],
            ["Développement modèle", "Entraînement, optimisation hyperparamètres, feature engineering", "Modèle entraîné, métriques"],
            ["Tests / validation", "Validation croisée, tests de robustesse, benchmark", "Rapport de tests, matrice de confusion"],
            ["Intégration", "API, pipeline de données, intégration SI existant", "API déployée, documentation technique"],
            ["Monitoring", "Surveillance des performances, détection de drift, alerting", "Dashboard monitoring, alertes"],
        ],
    )

    _add_sub_heading(doc, "3.2 Approche MLOps")
    _add_placeholder(
        doc,
        "[Décrire l'approche MLOps : automatisation des pipelines, "
        "gestion du cycle de vie des modèles, versioning des données "
        "et des modèles, CI/CD pour le ML, monitoring en production.]"
    )

    doc.add_page_break()

    # ── Section 4 — Gouvernance IA ──
    _add_section_heading(doc, 4, "Gouvernance IA")

    _add_sub_heading(doc, "4.1 Instances de pilotage")
    _make_grid_table(
        doc,
        ["Instance", "Fréquence", "Participants", "Objectif"],
        [
            [
                "Comité données",
                "Trimestriel",
                "Direction projet, DSI acheteur, DPO",
                "Stratégie données, conformité RGPD, arbitrages",
            ],
            [
                "Revue éthique IA",
                "À chaque livraison de modèle",
                "Responsable éthique, Data Scientist, référent acheteur",
                "Audit biais, transparence, conformité éthique",
            ],
            [
                "Point technique",
                "Hebdomadaire",
                "Équipe IA, référent technique acheteur",
                "Avancement développement, points bloquants",
            ],
            [
                "Revue performance modèle",
                "Mensuelle",
                "Data Scientist, ML Engineer, Product Owner",
                "Métriques de performance, drift, ré-entraînement",
            ],
        ],
    )

    _add_sub_heading(doc, "4.2 Reporting")
    reporting_items = [
        "Dashboard de performance des modèles en production",
        "Rapport de qualité des données",
        "Compte-rendu de revue éthique",
        "Plan d'amélioration continue",
    ]
    for item in reporting_items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        _set_run_font(run, size=10)

    doc.add_page_break()

    # ── Section 5 — Équipe IA ──
    _add_section_heading(doc, 5, "Équipe IA")

    _add_sub_heading(doc, "5.1 Profils mobilisés")
    _make_grid_table(
        doc,
        ["Rôle", "Nom", "Expérience", "Taux d'occupation"],
        [
            ["Chef de projet IA", "[Nom]", "[X ans]", "[X %]"],
            ["Data Scientist senior", "[Nom]", "[X ans]", "[X %]"],
            ["ML Engineer", "[Nom]", "[X ans]", "[X %]"],
            ["Data Engineer", "[Nom]", "[X ans]", "[X %]"],
            ["Développeur intégration", "[Nom]", "[X ans]", "[X %]"],
            ["Responsable éthique IA", "[Nom]", "[X ans]", "[X %]"],
        ],
    )

    _add_placeholder(
        doc,
        "[Les CV détaillés de chaque intervenant sont fournis en annexe.]"
    )

    doc.add_page_break()

    # ── Section 6 — Stack technique IA ──
    _add_section_heading(doc, 6, "Stack technique IA")

    _make_grid_table(
        doc,
        ["Domaine", "Technologies / Outils"],
        [
            ["Langages", "Python / R"],
            ["ML Frameworks", "TensorFlow / PyTorch / scikit-learn"],
            ["MLOps", "MLflow / Kubeflow"],
            ["Cloud", "AWS SageMaker / GCP Vertex / Azure ML"],
            ["Data", "Spark / Airflow"],
            ["API", "FastAPI / Flask"],
        ],
    )

    _add_placeholder(
        doc,
        "[Adapter la stack technique en fonction des exigences du CCTP "
        "et de l'environnement technique de l'acheteur.]"
    )

    doc.add_page_break()

    # ── Section 7 — Engagements performance IA ──
    _add_section_heading(doc, 7, "Engagements performance IA")

    _make_grid_table(
        doc,
        ["Indicateur", "Cible", "Mesure"],
        [
            ["Accuracy modèle", "> X %", "Métriques MLflow / rapport de tests"],
            ["Latence inférence", "< X ms", "Monitoring API (P95)"],
            ["Data quality score", "> 95 %", "Pipeline de validation données"],
            ["Drift detection", "Alerte automatique", "Monitoring modèle en production"],
        ],
    )

    _add_placeholder(
        doc,
        "[Adapter les indicateurs et cibles en fonction des exigences "
        "du CCAP et du CCTP. Préciser les seuils de performance attendus "
        "pour chaque modèle.]"
    )

    doc.add_page_break()

    # ── Section 8 — Réversibilité et éthique ──
    _add_section_heading(doc, 8, "Réversibilité et éthique")

    _add_sub_heading(doc, "8.1 Documentation des modèles")
    p = doc.add_paragraph()
    run = p.add_run(
        "Chaque modèle livré est accompagné d'une documentation complète : "
        "architecture du modèle, données d'entraînement, hyperparamètres, "
        "métriques de performance, limites connues. Le format Model Card "
        "est utilisé pour standardiser cette documentation."
    )
    _set_run_font(run, size=10)

    _add_sub_heading(doc, "8.2 Traçabilité des données (data lineage)")
    p = doc.add_paragraph()
    run = p.add_run(
        "Le lineage complet des données est maintenu : sources, "
        "transformations, versions utilisées pour l'entraînement. "
        "Cela garantit la reproductibilité et facilite les audits."
    )
    _set_run_font(run, size=10)

    _add_sub_heading(doc, "8.3 Conformité RGPD")
    p = doc.add_paragraph()
    run = p.add_run(
        "Le traitement des données personnelles est conforme au RGPD : "
        "minimisation des données, anonymisation/pseudonymisation, "
        "registre des traitements, analyse d'impact (PIA) si nécessaire."
    )
    _set_run_font(run, size=10)

    _add_sub_heading(doc, "8.4 Audit de biais")
    _add_placeholder(
        doc,
        "[Décrire la méthodologie d'audit de biais : métriques d'équité "
        "(demographic parity, equalized odds), outils utilisés (Fairlearn, "
        "AI Fairness 360), fréquence des audits, processus de correction.]"
    )

    _add_sub_heading(doc, "8.5 Transfert de ré-entraînement")
    _add_placeholder(
        doc,
        "[Décrire les modalités de transfert du processus de ré-entraînement "
        "des modèles : documentation des pipelines, formation de l'équipe "
        "interne, accès aux outils MLOps, procédures de validation.]"
    )

    doc.add_page_break()


# Dispatch map
_SECTION_BUILDERS: dict[str, callable] = {
    "TMA": _sections_tma,
    "DEVELOPPEMENT": _sections_dev,
    "FORMATION": _sections_formation,
    "IA": _sections_ia,
}


# ---------------------------------------------------------------------------
# Main generator
# ---------------------------------------------------------------------------

def generate_memoire(notice: dict[str, Any], dest_path: Path,
                     offer_type: str = "TMA") -> Path:
    """Generate a Mémoire Technique Word document for a BOAMP notice.

    Args:
        notice: BOAMP notice data dict.
        dest_path: Path to save the .docx file.
        offer_type: Type of offer — "TMA", "DEVELOPPEMENT", "FORMATION", or "IA".

    Returns:
        Path to the generated file.
    """
    offer_type = offer_type.upper()
    if offer_type not in _SECTION_BUILDERS:
        raise ValueError(
            f"Unknown offer_type {offer_type!r}. "
            f"Expected one of: {', '.join(_SECTION_BUILDERS)}"
        )

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)

    # Extract notice data
    objet = notice.get("objet", "[OBJET DU MARCHÉ]")
    acheteur = notice.get("nomacheteur", "[POUVOIR ADJUDICATEUR]")
    idweb = notice.get("idweb", "[REF]")

    # ══════════════════════════════════════════════════════════════════════
    # COVER PAGE
    # ══════════════════════════════════════════════════════════════════════
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    # Company name
    p_company = doc.add_paragraph()
    p_company.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_company.add_run(E["raison_sociale"])
    _set_run_font(run, bold=True, size=28, color=RGBColor(0x1F, 0x49, 0x7D))

    doc.add_paragraph()

    # Horizontal rule via single-row table
    rule_table = doc.add_table(rows=1, cols=1)
    rule_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_cell_shading(rule_table.cell(0, 0), "DAEEF3")
    rule_table.cell(0, 0).paragraphs[0].add_run(" ")

    doc.add_paragraph()

    # Document title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run("MÉMOIRE TECHNIQUE")
    _set_run_font(run, bold=True, size=22, color=RGBColor(0x1F, 0x49, 0x7D))

    doc.add_paragraph()

    # Market object
    p_obj = doc.add_paragraph()
    p_obj.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_obj.add_run(objet)
    _set_run_font(run, bold=True, size=14)

    doc.add_paragraph()

    # Buyer
    p_buyer = doc.add_paragraph()
    p_buyer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_buyer.add_run(f"Pouvoir adjudicateur : {acheteur}")
    _set_run_font(run, size=12)

    # BOAMP reference
    p_ref = doc.add_paragraph()
    p_ref.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_ref.add_run(f"Référence BOAMP : {idweb}")
    _set_run_font(run, size=12)

    doc.add_paragraph()
    doc.add_paragraph()

    # Contact info on cover
    p_contact = doc.add_paragraph()
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_contact.add_run(
        f"{E['adresse']} — {E['code_postal']} {E['ville']}\n"
        f"Tél. {E['telephone']} — {E['email']}"
    )
    _set_run_font(run, size=10, color=RGBColor(0x60, 0x60, 0x60))

    # Page break after cover
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # TABLE OF CONTENTS
    # ══════════════════════════════════════════════════════════════════════
    p_toc_title = doc.add_paragraph()
    run = p_toc_title.add_run("SOMMAIRE")
    _set_run_font(run, bold=True, size=16, color=RGBColor(0x1F, 0x49, 0x7D))
    p_toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    toc_items = _TOC_ITEMS[offer_type]

    for i, item in enumerate(toc_items, start=1):
        p = doc.add_paragraph()
        run = p.add_run(f"{i}.  {item}")
        _set_run_font(run, size=12)
        p.paragraph_format.space_after = Pt(4)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 1 — Présentation de l'entreprise
    # ══════════════════════════════════════════════════════════════════════
    _add_section_heading(doc, 1, "Présentation de l'entreprise")

    _add_sub_heading(doc, "1.1 Fiche d'identité")

    identity_rows = [
        ["Raison sociale", E["raison_sociale"]],
        ["Forme juridique", E["forme_juridique"]],
        ["SIRET", E["siret"]],
        ["N° TVA intracommunautaire", E["tva_intra"]],
        ["Capital social", E["capital_social"]],
        ["Date de création", E["date_creation"]],
        ["Code NAF", f"{E['code_naf']} — {E['libelle_naf']}"],
        ["Adresse", f"{E['adresse']}, {E['code_postal']} {E['ville']}"],
        ["Activité principale", E["description_activite"]],
        ["Dirigeant", f"{E['representant_prenom']} {E['representant_nom']}, {E['representant_qualite']}"],
        ["Convention collective", E["convention_collective"]],
        ["Catégorie", E["categorie"]],
    ]
    _make_grid_table(doc, ["Rubrique", "Détail"], identity_rows)

    _add_sub_heading(doc, "1.2 Positionnement et expertise")
    _add_placeholder(
        doc,
        "[Décrire le positionnement de Fenrir IT sur le marché des services "
        "numériques : domaines d'expertise (développement, TMA, infogérance, "
        "cloud), secteurs cibles (public, éducation, santé), avantages "
        "concurrentiels (proximité, réactivité, taille humaine).]"
    )

    _add_sub_heading(doc, "1.3 Certifications et labels")
    _add_placeholder(
        doc,
        "[Lister les certifications détenues ou en cours : ISO 27001, "
        "ITIL, PMP, labels qualité, agréments spécifiques au secteur public. "
        "Joindre les justificatifs en annexe.]"
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 2 — Compréhension du besoin
    # ══════════════════════════════════════════════════════════════════════
    _add_section_heading(doc, 2, "Compréhension du besoin")

    _add_sub_heading(doc, "2.1 Contexte et reformulation")
    _add_placeholder(
        doc,
        "[Reformuler le contexte du marché tel que compris à la lecture du "
        "CCTP et du règlement de consultation. Montrer la compréhension fine "
        "des enjeux de l'acheteur, de son environnement SI et de ses "
        "contraintes opérationnelles.]"
    )

    _add_sub_heading(doc, "2.2 Périmètre applicatif")
    _make_grid_table(
        doc,
        ["Application", "Technologie", "Criticité", "Observations"],
        [
            ["[Application 1]", "[Java/PHP/…]", "[Haute/Moyenne/Basse]", ""],
            ["[Application 2]", "[.NET/Python/…]", "[Haute/Moyenne/Basse]", ""],
            ["[Application 3]", "[…]", "[…]", ""],
        ],
    )

    _add_sub_heading(doc, "2.3 Enjeux identifiés")
    enjeux = [
        "Continuité de service et disponibilité des applications critiques",
        "Maîtrise des coûts de maintenance dans la durée",
        "Montée en compétence et transfert de connaissances",
        "Respect des contraintes réglementaires (RGPD, RGS, accessibilité)",
        "[Compléter avec les enjeux spécifiques identifiés dans le CCTP]",
    ]
    for enjeu in enjeux:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(enjeu)
        _set_run_font(run, size=10)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # SECTIONS 3-8 — Type-specific
    # ══════════════════════════════════════════════════════════════════════
    _SECTION_BUILDERS[offer_type](doc, notice)

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 9 — Références clients
    # ══════════════════════════════════════════════════════════════════════
    _add_section_heading(doc, 9, "Références clients")

    for ref_num in range(1, 4):
        _add_sub_heading(doc, f"Référence {ref_num}")
        _make_grid_table(
            doc,
            ["Rubrique", "Détail"],
            [
                ["Client", "[Nom du client]"],
                ["Secteur", "[Public / Éducation / Santé / Privé]"],
                ["Objet de la prestation", "[Description de la TMA réalisée]"],
                ["Durée", "[Date début — Date fin]"],
                ["Montant annuel", "[Montant HT]"],
                ["Environnement technique", "[Technologies, outils]"],
                ["Contact référent", "[Nom, fonction, téléphone, email]"],
            ],
        )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 10 — Annexes
    # ══════════════════════════════════════════════════════════════════════
    _add_section_heading(doc, 10, "Annexes")

    p = doc.add_paragraph()
    run = p.add_run("Checklist des pièces jointes :")
    _set_run_font(run, bold=True, size=11)

    annexes = [
        "CV détaillés des intervenants proposés",
        "Attestations d'assurance responsabilité civile professionnelle",
        "Certifications (ISO, ITIL, etc.)",
        "Exemples de livrables (tableau de bord, PV de recette)",
        "Attestations de références clients",
        "Plan qualité type",
    ]
    for annexe in annexes:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"☐  {annexe}")
        _set_run_font(run, size=10)

    doc.add_paragraph()

    # Footer
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_p.add_run(
        f"Mémoire Technique — {E['raison_sociale']} — {idweb}"
    )
    _set_run_font(footer_run, size=8, italic=True,
                  color=RGBColor(0x80, 0x80, 0x80))

    # ── Save ──
    dest_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(dest_path))
    return dest_path
