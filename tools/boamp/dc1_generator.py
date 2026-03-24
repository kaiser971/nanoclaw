"""Génération du formulaire DC1 au format Word (.docx).

Produit un DC1 conforme au modèle officiel du Ministère de l'Économie,
pré-rempli avec les informations Fenrir IT et les données de l'offre BOAMP.
"""

from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

from entreprise import ENTREPRISE as E


def _set_cell_shading(cell, color: str) -> None:
    """Set cell background color."""
    from docx.oxml.ns import qn
    from lxml import etree

    shading = etree.SubElement(
        cell._element.get_or_add_tcPr(), qn("w:shd")
    )
    shading.set(qn("w:fill"), color)
    shading.set(qn("w:val"), "clear")


def _add_heading_cell(table, row: int, col: int, text: str) -> None:
    """Add a blue-shaded heading cell."""
    cell = table.cell(row, col)
    cell.text = text
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(11)
    _set_cell_shading(cell, "DAEEF3")


def _add_section_heading(doc: Document, text: str) -> None:
    """Add a blue-background section heading like the official DC1."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = True
    run.underline = True
    run.font.size = Pt(11)
    _set_cell_shading(cell, "DAEEF3")


def generate_dc1(notice: dict[str, Any], dest_path: Path) -> Path:
    """Generate a DC1 Word document for a BOAMP notice.

    Args:
        notice: BOAMP notice data dict.
        dest_path: Path to save the .docx file.

    Returns:
        Path to the generated file.
    """
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)

    # Extract notice data
    objet = notice.get("objet", "[OBJET DU MARCHÉ]")
    acheteur = notice.get("nomacheteur", "[POUVOIR ADJUDICATEUR]")
    idweb = notice.get("idweb", "[REF]")
    date_limite = notice.get("datelimitereponse", "")
    nature = notice.get("nature_categorise_libelle", "")
    procedure = notice.get("procedure_categorise", "")
    type_marche = notice.get("type_marche", "SERVICES")

    # ── Header ──
    header_p = doc.add_paragraph()
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_p.add_run("MINISTÈRE DE L'ÉCONOMIE ET DES FINANCES")
    run.font.size = Pt(9)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sub.add_run("Direction des Affaires Juridiques")
    run2.font.size = Pt(9)
    run2.italic = True

    doc.add_paragraph()

    # Title block
    title_table = doc.add_table(rows=3, cols=2)
    title_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    cell_left = title_table.cell(0, 0)
    p = cell_left.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("MARCHÉS PUBLICS ET ACCORDS-CADRES")
    run.bold = True
    run.font.size = Pt(12)

    cell_right = title_table.cell(0, 1)
    p = cell_right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DC1")
    run.bold = True
    run.font.size = Pt(16)

    title_table.cell(0, 0).merge(title_table.cell(0, 0))

    title2 = title_table.cell(1, 0)
    title2.merge(title_table.cell(1, 1))
    p = title2.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("LETTRE DE CANDIDATURE")
    run.bold = True
    run.font.size = Pt(14)

    title3 = title_table.cell(2, 0)
    title3.merge(title_table.cell(2, 1))
    p = title3.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("HABILITATION DU MANDATAIRE PAR SES CO-TRAITANTS")
    run.bold = True
    run.font.size = Pt(11)

    doc.add_paragraph()

    # Intro text
    intro = doc.add_paragraph()
    intro_run = intro.add_run(
        "Le formulaire DC1 est un modèle de lettre de candidature, "
        "qui peut être utilisé par les candidats aux marchés publics ou "
        "accords-cadres pour présenter leur candidature. "
        "En cas d'allotissement, ce document peut être commun à plusieurs lots."
    )
    intro_run.font.size = Pt(8)
    intro_run.italic = True

    doc.add_paragraph()

    # ── A - Identification du pouvoir adjudicateur ──
    _add_section_heading(doc, "A - Identification du pouvoir adjudicateur (ou de l'entité adjudicatrice).")

    a_text = doc.add_paragraph()
    a_text.add_run(acheteur).bold = True
    doc.add_paragraph(f"Référence BOAMP : {idweb}")
    if date_limite:
        dl_str = date_limite[:10] if isinstance(date_limite, str) else str(date_limite)
        doc.add_paragraph(f"Date limite de réponse : {dl_str}")

    doc.add_paragraph()

    # ── B - Objet de la consultation ──
    _add_section_heading(doc, "B - Objet de la consultation.")

    b_text = doc.add_paragraph()
    run = b_text.add_run(objet)
    run.bold = True

    b_detail = doc.add_paragraph()
    details = []
    if type_marche:
        details.append(f"Type de marché : {type_marche}")
    if procedure:
        details.append(f"Procédure : {procedure}")
    if nature:
        details.append(f"Nature : {nature}")
    b_detail.add_run(" — ".join(details)).font.size = Pt(9)

    doc.add_paragraph()

    # ── C - Objet de la candidature ──
    _add_section_heading(doc, "C - Objet de la candidature.")

    c_intro = doc.add_paragraph("La candidature est présentée :")
    c_intro.paragraph_format.space_after = Pt(4)

    c1 = doc.add_paragraph(style="List Bullet")
    run = c1.add_run("☒ pour le marché public ou pour l'accord-cadre")
    run.font.size = Pt(10)

    doc.add_paragraph()

    # ── D - Présentation du candidat ──
    _add_section_heading(doc, "D - Présentation du candidat.")

    d_intro = doc.add_paragraph()
    d_intro.add_run("☒ Le candidat se présente seul :").bold = True

    doc.add_paragraph()

    # Candidate info table
    info_table = doc.add_table(rows=8, cols=2)
    info_table.style = "Table Grid"
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    fields = [
        ("Nom commercial et dénomination sociale", E["raison_sociale"]),
        ("Forme juridique", E["forme_juridique"]),
        (
            "Adresse du siège social",
            f"{E['adresse']}\n{E['code_postal']} {E['ville']}",
        ),
        ("Téléphone", E["telephone"]),
        ("Adresse électronique", E["email"]),
        (
            "Personne(s) ayant le pouvoir d'engager la société",
            f"{E['representant_prenom']} {E['representant_nom']}, {E['representant_qualite']}",
        ),
        ("Numéro SIRET", E["siret"]),
        ("Code APE", f"{E['code_naf']} — {E['libelle_naf']}"),
    ]

    for i, (label, value) in enumerate(fields):
        _set_cell_shading(info_table.cell(i, 0), "F2F2F2")
        label_p = info_table.cell(i, 0).paragraphs[0]
        label_run = label_p.add_run(label)
        label_run.bold = True
        label_run.font.size = Pt(9)

        val_p = info_table.cell(i, 1).paragraphs[0]
        val_run = val_p.add_run(value)
        val_run.font.size = Pt(10)

    # Set column widths
    for row in info_table.rows:
        row.cells[0].width = Cm(6)
        row.cells[1].width = Cm(11)

    doc.add_paragraph()

    d_group = doc.add_paragraph()
    d_group.add_run("☐ Le candidat est un groupement d'entreprises").font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    d_group.add_run(" (non applicable)").font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.add_paragraph()

    # ── E - Identification des membres du groupement ──
    _add_section_heading(doc, "E - Identification des membres du groupement et répartition des prestations.")

    e_text = doc.add_paragraph()
    run = e_text.add_run("Sans objet — candidature individuelle.")
    run.italic = True
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.add_paragraph()

    # ── F - Engagements ──
    _add_section_heading(doc, "F - Engagements du candidat individuel ou de chaque membre du groupement.")

    f1_title = doc.add_paragraph()
    f1_title.add_run("F1 - Attestations sur l'honneur").bold = True

    f1_intro = doc.add_paragraph(
        "Le candidat individuel déclare sur l'honneur :"
    )
    f1_intro.paragraph_format.space_after = Pt(4)

    attestations = [
        "a) Condamnation définitive : ne pas avoir fait l'objet, depuis moins de cinq ans, "
        "d'une condamnation définitive pour l'une des infractions prévues aux articles 222-38, "
        "222-40, 226-13, 313-1 à 313-3, 314-1 à 314-3, 324-1 à 324-6, 413-9 à 413-12, "
        "421-1 à 421-2-3 du code pénal ;",
        "b) Lutte contre le travail illégal : ne pas avoir fait l'objet, depuis moins de cinq ans, "
        "d'une condamnation inscrite au bulletin n° 2 du casier judiciaire pour les infractions "
        "mentionnées aux articles L. 8221-1, L. 8221-3, L. 8221-5, L. 8231-1, L. 8241-1, "
        "L. 8251-1 et L. 8251-2 du code du travail ;",
        "c) Obligation d'emploi des travailleurs handicapés : être en règle au regard "
        "des articles L. 5212-1 à L. 5212-11 du code du travail ;",
        "d) Liquidation judiciaire : ne pas être soumis à la procédure de liquidation judiciaire ;",
        "e) Redressement judiciaire : ne pas être admis à la procédure de redressement judiciaire ;",
        "f) Situation fiscale et sociale : avoir souscrit les déclarations lui incombant en "
        "matière fiscale et sociale et acquitté les impôts et cotisations exigibles ;",
        "h) que les renseignements fournis dans le formulaire DC2 sont exacts.",
    ]

    for att in attestations:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"☒ {att}")
        run.font.size = Pt(8)

    doc.add_paragraph()

    f2_title = doc.add_paragraph()
    f2_title.add_run("F2 - Capacités.").bold = True

    f2_text = doc.add_paragraph(
        "Le candidat déclare présenter les capacités nécessaires à l'exécution "
        "du marché public et produit à cet effet :"
    )
    f2_check = doc.add_paragraph()
    f2_check.add_run("☒ le formulaire DC2.").bold = True

    doc.add_paragraph()

    # ── G - Désignation du mandataire ──
    _add_section_heading(doc, "G - Désignation du mandataire et habilitation (en cas de groupement).")

    g_text = doc.add_paragraph()
    run = g_text.add_run("Sans objet — candidature individuelle.")
    run.italic = True
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.add_paragraph()

    # ── H - Signature ──
    _add_section_heading(doc, "H - Signature du candidat individuel ou de chaque membre du groupement.")

    doc.add_paragraph()

    sig_table = doc.add_table(rows=2, cols=3)
    sig_table.style = "Table Grid"
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["Nom, prénom et qualité\ndu signataire", "Lieu et date de signature", "Signature"]
    for i, h in enumerate(headers):
        _set_cell_shading(sig_table.cell(0, i), "DAEEF3")
        p = sig_table.cell(0, i).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)

    # Pre-fill signature row
    sig_table.cell(1, 0).paragraphs[0].add_run(
        f"{E['representant_prenom']} {E['representant_nom']}\n{E['representant_qualite']}"
    ).font.size = Pt(10)

    sig_table.cell(1, 1).paragraphs[0].add_run(
        f"{E['ville']}, le ____/____/________"
    ).font.size = Pt(10)

    sig_table.cell(1, 2).paragraphs[0].add_run(
        "\n\n\n"
    )

    # Set signature row height for space
    for cell in sig_table.rows[1].cells:
        cell.height = Cm(3)

    doc.add_paragraph()

    # Footer
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_p.add_run(
        f"DC1 – Lettre de candidature          {idweb}          "
        f"Page 1 / 1"
    )
    footer_run.font.size = Pt(8)
    footer_run.italic = True

    # Save
    dest_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(dest_path))
    return dest_path
