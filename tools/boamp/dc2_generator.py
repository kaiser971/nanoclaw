"""Génération du formulaire DC2 au format Word (.docx).

Produit un DC2 conforme au modèle officiel du Ministère de l'Économie,
pré-rempli avec les informations Fenrir IT et les données de l'offre BOAMP.
"""

from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

from entreprise import ENTREPRISE as E


# ── Helper functions (same pattern as dc1_generator.py) ──────────────────────


def _set_cell_shading(cell, color: str) -> None:
    """Set cell background color."""
    from docx.oxml.ns import qn
    from lxml import etree

    shading = etree.SubElement(
        cell._element.get_or_add_tcPr(), qn("w:shd")
    )
    shading.set(qn("w:fill"), color)
    shading.set(qn("w:val"), "clear")


def _add_heading_cell(table, row: int, col: int, text: str) -> None:
    """Add a blue-shaded heading cell."""
    cell = table.cell(row, col)
    cell.text = text
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(11)
    _set_cell_shading(cell, "DAEEF3")


def _add_section_heading(doc: Document, text: str) -> None:
    """Add a blue-background section heading like the official DC2."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = True
    run.underline = True
    run.font.size = Pt(11)
    _set_cell_shading(cell, "DAEEF3")


def _add_greyed_text(doc: Document, text: str) -> None:
    """Add a paragraph with greyed-out italic text."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


# ── Main generator ───────────────────────────────────────────────────────────


def generate_dc2(notice: dict[str, Any], dest_path: Path) -> Path:
    """Generate a DC2 Word document for a BOAMP notice.

    Args:
        notice: BOAMP notice data dict.
        dest_path: Path to save the .docx file.

    Returns:
        Path to the generated file.
    """
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)

    # Extract notice data
    objet = notice.get("objet", "[OBJET DU MARCHÉ]")
    acheteur = notice.get("nomacheteur", "[POUVOIR ADJUDICATEUR]")
    idweb = notice.get("idweb", "[REF]")
    nature = notice.get("nature_categorise_libelle", "")
    procedure = notice.get("procedure_categorise", "")
    type_marche = notice.get("type_marche", "SERVICES")

    # ── Header ──
    header_p = doc.add_paragraph()
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_p.add_run("MINISTÈRE DE L'ÉCONOMIE ET DES FINANCES")
    run.font.size = Pt(9)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sub.add_run("Direction des Affaires Juridiques")
    run2.font.size = Pt(9)
    run2.italic = True

    doc.add_paragraph()

    # ── Title block ──
    title_table = doc.add_table(rows=3, cols=2)
    title_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    cell_left = title_table.cell(0, 0)
    p = cell_left.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("MARCHÉS PUBLICS ET ACCORDS-CADRES")
    run.bold = True
    run.font.size = Pt(12)

    cell_right = title_table.cell(0, 1)
    p = cell_right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DC2")
    run.bold = True
    run.font.size = Pt(16)

    title2 = title_table.cell(1, 0)
    title2.merge(title_table.cell(1, 1))
    p = title2.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DÉCLARATION DU CANDIDAT INDIVIDUEL")
    run.bold = True
    run.font.size = Pt(14)

    title3 = title_table.cell(2, 0)
    title3.merge(title_table.cell(2, 1))
    p = title3.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("OU DU MEMBRE DU GROUPEMENT")
    run.bold = True
    run.font.size = Pt(11)

    doc.add_paragraph()

    # Intro text
    intro = doc.add_paragraph()
    intro_run = intro.add_run(
        "Le formulaire DC2 est un modèle de déclaration qui peut être utilisé "
        "par les candidats aux marchés publics ou accords-cadres à l'appui de "
        "leur candidature (formulaire DC1)."
    )
    intro_run.font.size = Pt(8)
    intro_run.italic = True

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # Section A - Identification du pouvoir adjudicateur
    # ══════════════════════════════════════════════════════════════════════════
    _add_section_heading(
        doc, "A - Identification du pouvoir adjudicateur (ou de l'entité adjudicatrice)."
    )

    a_text = doc.add_paragraph()
    a_text.add_run(acheteur).bold = True
    doc.add_paragraph(f"Référence BOAMP : {idweb}")

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # Section B - Objet du marché
    # ══════════════════════════════════════════════════════════════════════════
    _add_section_heading(doc, "B - Objet du marché public.")

    b_text = doc.add_paragraph()
    run = b_text.add_run(objet)
    run.bold = True

    b_detail = doc.add_paragraph()
    details = []
    if type_marche:
        details.append(f"Type de marché : {type_marche}")
    if procedure:
        details.append(f"Procédure : {procedure}")
    if nature:
        details.append(f"Nature : {nature}")
    b_detail.add_run(" — ".join(details)).font.size = Pt(9)

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # Section C - Identification du candidat
    # ══════════════════════════════════════════════════════════════════════════
    _add_section_heading(doc, "C - Identification du candidat individuel ou du membre du groupement.")

    # C1 - Cas général
    c1_title = doc.add_paragraph()
    c1_title.add_run("C1 - Cas général :").bold = True

    doc.add_paragraph()

    c1_table = doc.add_table(rows=8, cols=2)
    c1_table.style = "Table Grid"
    c1_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    c1_fields = [
        ("Nom commercial et dénomination sociale", E["raison_sociale"]),
        (
            "Adresse",
            f"{E['adresse']}\n{E['code_postal']} {E['ville']}",
        ),
        ("Adresse électronique", E["email"]),
        ("Téléphone", E["telephone"]),
        ("Numéro SIRET", E["siret"]),
        ("Forme juridique", E["forme_juridique"]),
        (
            "Personne(s) ayant le pouvoir d'engager le candidat",
            f"{E['representant_prenom']} {E['representant_nom']}, {E['representant_qualite']}",
        ),
        ("Code APE", f"{E['code_naf']} — {E['libelle_naf']}"),
    ]

    for i, (label, value) in enumerate(c1_fields):
        _set_cell_shading(c1_table.cell(i, 0), "F2F2F2")
        label_p = c1_table.cell(i, 0).paragraphs[0]
        label_run = label_p.add_run(label)
        label_run.bold = True
        label_run.font.size = Pt(9)

        val_p = c1_table.cell(i, 1).paragraphs[0]
        val_run = val_p.add_run(value)
        val_run.font.size = Pt(10)

    for row in c1_table.rows:
        row.cells[0].width = Cm(6)
        row.cells[1].width = Cm(11)

    doc.add_paragraph()

    # C2 - Cas particuliers
    c2_title = doc.add_paragraph()
    c2_run = c2_title.add_run("C2 - Cas particuliers :")
    c2_run.bold = True
    c2_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    _add_greyed_text(doc, "Sans objet")

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # Section D - Renseignements financiers
    # ══════════════════════════════════════════════════════════════════════════
    _add_section_heading(doc, "D - Renseignements d'ordre juridique, économique, financier et technique.")

    # D1 - Chiffre d'affaires
    d1_title = doc.add_paragraph()
    d1_title.add_run("D1 - Chiffre d'affaires :").bold = True

    doc.add_paragraph()

    ca_table = doc.add_table(rows=4, cols=3)
    ca_table.style = "Table Grid"
    ca_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    ca_headers = ["Année", "CA global (€ HT)", "CA lié au domaine d'activité\nconcerné par le marché (€ HT)"]
    for i, h in enumerate(ca_headers):
        _set_cell_shading(ca_table.cell(0, i), "DAEEF3")
        p = ca_table.cell(0, i).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)

    # Year rows
    years = ["2023", "2024", "2025"]
    for row_idx, year in enumerate(years, start=1):
        p = ca_table.cell(row_idx, 0).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(year)
        run.font.size = Pt(10)
        # Leave CA cells empty for manual fill
        ca_table.cell(row_idx, 1).paragraphs[0].add_run("").font.size = Pt(10)
        ca_table.cell(row_idx, 2).paragraphs[0].add_run("").font.size = Pt(10)

    for row in ca_table.rows:
        row.cells[0].width = Cm(3)
        row.cells[1].width = Cm(5)
        row.cells[2].width = Cm(7)

    doc.add_paragraph()

    # D2 - Redressement judiciaire
    d2_title = doc.add_paragraph()
    d2_title.add_run("D2 - Le candidat est-il en redressement judiciaire ?").bold = True

    d2_answer = doc.add_paragraph()
    run = d2_answer.add_run("\u2612 NON")
    run.font.size = Pt(11)
    run.bold = True

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # Section E - Capacités des opérateurs économiques
    # ══════════════════════════════════════════════════════════════════════════
    _add_section_heading(
        doc,
        "E - Capacités professionnelles, techniques et financières "
        "des autres opérateurs économiques sur lesquelles le candidat s'appuie.",
    )

    _add_greyed_text(doc, "Sans objet \u2014 candidature individuelle.")

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # Section F - Nationalité
    # ══════════════════════════════════════════════════════════════════════════
    _add_section_heading(doc, "F - Nationalité du candidat.")

    f_text = doc.add_paragraph()
    run = f_text.add_run("Française")
    run.font.size = Pt(10)

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # Section G - Récapitulatif des pièces
    # ══════════════════════════════════════════════════════════════════════════
    _add_section_heading(doc, "G - Récapitulatif des pièces demandées par le pouvoir adjudicateur.")

    g_intro = doc.add_paragraph()
    g_intro.add_run("Documents joints à la candidature :").bold = True

    doc.add_paragraph()

    pieces = [
        "Extrait Kbis de moins de 3 mois",
        "Attestation de régularité fiscale (DGFIP)",
        "Attestation de régularité sociale (URSSAF)",
        "Attestation d'assurance responsabilité civile professionnelle (RC Pro) en cours de validité",
        "Bilans et comptes de résultat des trois derniers exercices",
        "Déclaration sur le chiffre d'affaires des trois derniers exercices",
        "Certificats de capacité / références clients",
        "Formulaire DC1 \u2014 Lettre de candidature",
        "Formulaire DC2 \u2014 Déclaration du candidat (le présent document)",
    ]

    for piece in pieces:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"\u2612 {piece}")
        run.font.size = Pt(9)

    doc.add_paragraph()

    # ── Footer ──
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_p.add_run(
        f"DC2 \u2013 D\u00e9claration du candidat          {idweb}          "
        f"Page 1 / 1"
    )
    footer_run.font.size = Pt(8)
    footer_run.italic = True

    # Save
    dest_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(dest_path))
    return dest_path
